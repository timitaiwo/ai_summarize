""" """

import os
import re
import sys
import threading

from dataclasses import dataclass


import logging

import httpx

from docx import Document

from openai import OpenAI







# ai_client: OpenAI = GEMINI_CLIENT,
# model_name: str = "gemini-2.0-flash",
# system_message: str = "You are an admissions officer",








def generate_docx(
    essay_list: list[AuthorEssayKV], document_name: str = "document2.docx"
):
    ms_document = Document()

    for essay in essay_list:
        ms_document.add_heading("Author Details")
        ms_document.add_paragraph(essay["author_details"])

        ms_document.add_heading("Author Sample Essay(s)")
        ms_document.add_paragraph(essay["author_essay"])

    ms_document.save(document_name)


def main(base_url: str):
    valid_urls: set[str] = extract_valid_essay_urls(base_url)

    print("Escalabe")

    if len(valid_urls) == 0:
        print("No valid URLs detected")
        exit(1)

    # Threaded
    print("Starting threaded")
    ThreadedAuthorEssayKVList = extract_essays_from_urls(
        valid_urls=valid_urls, threaded=True
    )

    # Sequential
    SeqAuthorEssayKVList = extract_essays_from_urls(valid_urls=valid_urls)

    print(SeqAuthorEssayKVList == ThreadedAuthorEssayKVList)

    # CreateDocx
    generate_docx(SeqAuthorEssayKVList, "sequential_document.docx")
    generate_docx(ThreadedAuthorEssayKVList, "threaded_document.docx")
